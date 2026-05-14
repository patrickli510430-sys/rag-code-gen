from __future__ import annotations

import logging
from typing import Any

from fastapi import APIRouter, HTTPException, UploadFile, File, Form

from src.api.schemas import (
    BenchmarkRequest,
    CrawlRequest,
    CodeExecuteRequest,
    CodeValidateRequest,
    DocumentIndexRequest,
    GenerationResult,
    MetricsSummary,
    PipelineAskRequest,
    PythonCompleteRequest,
    PythonGenerateRequest,
    SandboxExecutionResult,
    SearchRequest,
    SQLCompleteRequest,
    SQLGenerateRequest,
    ValidationResult,
)
from src.crawler.code_crawler import CodeCrawler
from src.crawler.doc_crawler import DocCrawler
from src.evaluator.benchmark import BenchmarkRunner
from src.evaluator.metrics import get_metrics_tracker
from src.llm.python_generator import PythonGenerator
from src.llm.sql_generator import SQLGenerator
from src.rag.retriever import Retriever
from src.validator.python_validator import PythonValidator
from src.validator.sandbox import CodeSandbox
from src.validator.sql_validator import SQLValidator

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/api/v1", tags=["code-gen"])

_retriever = Retriever()
_sql_generator = SQLGenerator(retriever=_retriever)
_python_generator = PythonGenerator(retriever=_retriever)
_sql_validator = SQLValidator()
_python_validator = PythonValidator()
_sandbox = CodeSandbox()


# ---------------------------------------------------------------------------
#  LLM response normalization helpers
# ---------------------------------------------------------------------------
def _find_items_array(data: dict) -> list[dict]:
    item_keys = ("items", "review_results", "review_items", "findings")
    for key in item_keys:
        if key in data and isinstance(data[key], list):
            return data[key]
    for value in data.values():
        if isinstance(value, dict):
            for key in item_keys:
                if key in value and isinstance(value[key], list):
                    return value[key]
        elif isinstance(value, list) and value and isinstance(value[0], dict):
            if any(k in value[0] for k in ("clause", "regulation", "rule", "status", "compliance_status")):
                return value
    return []


def _find_value_recursive(data: dict, target_key: str) -> str | None:
    if target_key in data and data[target_key]:
        return str(data[target_key])
    for value in data.values():
        if isinstance(value, dict):
            result = _find_value_recursive(value, target_key)
            if result:
                return result
    return None


# ---------------------------------------------------------------------------
#  Collections
# ---------------------------------------------------------------------------
KNOWN_COLLECTIONS = {
    "regulatory_docs": "制度文档库 — 银行监管政策、法规、制度文件",
    "code_docs": "代码文档库 — SQL/Python 代码示例与技术文档",
}


def _get_retriever_for(collection: str = "regulatory_docs") -> Retriever:
    """Get or create a Retriever for a specific collection."""
    return Retriever(collection_name=collection)


# ===================================================================
#  Code Generation
# ===================================================================

@router.post("/sql/generate", response_model=GenerationResult)
async def generate_sql(request: SQLGenerateRequest):
    try:
        result = await _sql_generator.generate(
            requirement=request.requirement,
            table_schema=request.table_schema,
            use_rag=request.use_rag,
        )
        return GenerationResult(
            sql=result["sql"],
            raw_response=result["raw_response"],
            model=result["model"],
            usage=result["usage"],
            generation_time_ms=result["generation_time_ms"],
            context_used=result["context_used"],
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/python/generate", response_model=GenerationResult)
async def generate_python(request: PythonGenerateRequest):
    try:
        result = await _python_generator.generate(
            requirement=request.requirement,
            additional_context=request.additional_context,
            use_rag=request.use_rag,
        )
        return GenerationResult(
            code=result["code"],
            raw_response=result["raw_response"],
            model=result["model"],
            usage=result["usage"],
            generation_time_ms=result["generation_time_ms"],
            context_used=result["context_used"],
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/validate", response_model=ValidationResult)
async def validate_code(request: CodeValidateRequest):
    if request.language == "sql":
        result = _sql_validator.validate(request.code)
    elif request.language == "python":
        result = _python_validator.validate(request.code)
    else:
        raise HTTPException(status_code=400, detail=f"Unsupported language: {request.language}")
    return ValidationResult(
        is_valid=result.is_valid,
        is_safe=result.is_safe,
        errors=result.errors,
        warnings=result.warnings,
    )


@router.post("/execute", response_model=SandboxExecutionResult)
async def execute_code(request: CodeExecuteRequest):
    if request.language == "sql":
        result = _sandbox.execute_sql(request.code, table_schema=request.table_schema)
    elif request.language == "python":
        result = _sandbox.execute_python(request.code)
    else:
        raise HTTPException(status_code=400, detail=f"Unsupported language: {request.language}")
    return SandboxExecutionResult(
        success=result.success,
        output=result.output,
        error=result.error,
        exit_code=result.exit_code,
        execution_time_ms=result.execution_time_ms,
    )


@router.post("/pipeline/ask")
async def ask_for_code(request: PipelineAskRequest):
    if request.language == "sql":
        result = await _sql_generator.generate(
            requirement=request.requirement,
            table_schema=request.table_schema or "",
            use_rag=True,
        )
        generated = result["sql"]
    elif request.language == "python":
        result = await _python_generator.generate(
            requirement=request.requirement,
            additional_context=request.table_schema or "",
            use_rag=True,
        )
        generated = result["code"]
    else:
        raise HTTPException(status_code=400, detail=f"Unsupported language: {request.language}")

    search_results = _retriever.retrieve(request.requirement, top_k=3)

    if request.language == "sql":
        validation = _sql_validator.validate(generated)
    else:
        validation = _python_validator.validate(generated)

    # Sandbox: skip execution if no schema for SQL (tables won't exist)
    if request.language == "sql" and (not request.table_schema or not request.table_schema.strip()):
        sandbox_response = {
            "success": True,
            "output": "",
            "error": "",
            "note": "当前为通用模式（无业务表），沙箱未用生成的 SQL 执行。请选择银行场景后再试。",
            "skipped": True,
        }
    elif request.language == "sql":
        sandbox_result = _sandbox.execute_sql(generated, table_schema=request.table_schema)
        sandbox_response = {
            "success": sandbox_result.success,
            "output": sandbox_result.output[:500] if sandbox_result.output else "",
            "error": sandbox_result.error[:500] if sandbox_result.error else "",
        }
    else:
        sandbox_result = _sandbox.execute_python(generated)
        sandbox_response = {
            "success": sandbox_result.success,
            "output": sandbox_result.output[:500] if sandbox_result.output else "",
            "error": sandbox_result.error[:500] if sandbox_result.error else "",
        }

    response_data = {
        "requirement": request.requirement,
        "language": request.language,
        "generated_code": generated,
        "model": result["model"],
        "generation_time_ms": result["generation_time_ms"],
        "tokens_used": result["usage"].get("total_tokens", 0),
        "retrieved_context": [
            {"document": r.document[:200], "score": round(r.score, 4), "source": r.metadata.get("url", "")}
            for r in search_results[:3]
        ],
        "validation": {
            "is_valid": validation.is_valid,
            "is_safe": validation.is_safe,
            "errors": validation.errors,
        },
        "sandbox_execution": sandbox_response,
    }

    get_metrics_tracker().record_generation(
        gen_type=request.language,
        success=validation.is_valid and validation.is_safe,
        generation_time_ms=result["generation_time_ms"],
        validation_passed=validation.is_valid,
        rag_used=bool(search_results),
        tokens_used=response_data["tokens_used"],
        metadata={"requirement": request.requirement[:100]},
    )

    return response_data


# ===================================================================
#  Collections API
# ===================================================================

@router.get("/collections")
async def list_collections():
    """列出所有知识库集合"""
    result = {}
    for name, desc in KNOWN_COLLECTIONS.items():
        r = Retriever(collection_name=name)
        result[name] = {"name": name, "description": desc, "count": r.vector_store.count()}
    return result


# ===================================================================
#  Document Management (with collection support)
# ===================================================================

def _get_collection(request_data: dict | None = None, default: str = "regulatory_docs") -> str:
    """Extract collection name from request body or fall back to default."""
    if request_data and isinstance(request_data, dict):
        coll = request_data.get("collection", "")
        if coll in KNOWN_COLLECTIONS:
            return coll
    return default


@router.post("/documents/index")
async def index_documents(request: DocumentIndexRequest):
    collection = _get_collection(request.documents[0].get("metadata", {}) if request.documents else {}, "regulatory_docs")
    ret = _get_retriever_for(collection)
    count = ret.index_documents(request.documents)
    return {"indexed_count": count, "total_in_store": ret.vector_store.count(), "collection": collection}


@router.get("/documents")
async def list_documents(limit: int = 50, offset: int = 0, collection: str = "regulatory_docs"):
    ret = _get_retriever_for(collection)
    result = ret.vector_store.list_all(limit=limit, offset=offset)
    result["collection"] = collection
    return result


@router.delete("/documents")
async def clear_documents(collection: str = "regulatory_docs"):
    ret = _get_retriever_for(collection)
    count = ret.vector_store.count()
    ret.vector_store.delete_collection()
    return {"deleted": True, "previous_count": count, "collection": collection}


@router.delete("/documents/by-id")
async def delete_document_by_body(data: dict):
    doc_id = data.get("id", "")
    if not doc_id:
        raise HTTPException(status_code=400, detail="Missing 'id' in request body")
    collection = _get_collection(data, "regulatory_docs")
    ret = _get_retriever_for(collection)
    deleted = ret.vector_store.delete_by_ids([doc_id])
    return {"deleted": deleted, "id": doc_id, "collection": collection}


# ===================================================================
#  File Upload (drag-and-drop local documents)
# ===================================================================

@router.post("/documents/upload")
async def upload_file(
    file: UploadFile = File(...),
    collection: str = Form(default="regulatory_docs"),
):
    """上传本地文档文件（支持 .txt .docx .md），自动索引入库"""
    import re

    filename = file.filename or "unknown"
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    SUPPORTED_UPLOAD = {"txt", "md", "docx", "csv", "json", "py", "sql", "js", "ts", "java", "go", "sh", "cpp", "h", "html", "css", "yaml", "yml", "toml", "ini", "cfg"}
    if ext not in SUPPORTED_UPLOAD:
        raise HTTPException(status_code=400, detail=f"不支持的文件格式: .{ext}（支持代码文件和文档文件）")

    try:
        raw = await file.read()
    except Exception:
        raise HTTPException(status_code=400, detail="读取文件失败")

    content = ""
    if ext == "docx":
        try:
            from io import BytesIO
            from docx import Document
            doc = Document(BytesIO(raw))
            content = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        except ImportError:
            raise HTTPException(status_code=500, detail="服务器缺少 python-docx 库，无法解析 .docx 文件")
        except Exception as e:
            raise HTTPException(status_code=400, detail=f".docx 解析失败: {str(e)[:200]}")
    else:
        # txt, md, csv, json — plain text
        try:
            content = raw.decode("utf-8")
        except UnicodeDecodeError:
            try:
                content = raw.decode("gbk")
            except Exception:
                raise HTTPException(status_code=400, detail="文件编码无法识别，请转为 UTF-8")

    if not content or not content.strip():
        raise HTTPException(status_code=400, detail="文件内容为空")

    ret = _get_retriever_for(collection)
    count = ret.index_documents([{
        "content": content,
        "url": f"file://{filename}",
        "metadata": {"source": "upload", "filename": filename, "file_type": ext},
    }])

    return {
        "filename": filename,
        "file_type": ext,
        "indexed_count": count,
        "total_in_store": ret.vector_store.count(),
        "collection": collection,
        "preview": content[:200],
    }


@router.post("/documents/import-folder")
async def import_folder(data: dict):
    """从本地文件夹批量导入文档。"""
    import os as _os
    folder_path = data.get("path", "")
    collection = data.get("collection", "regulatory_docs")
    if collection not in KNOWN_COLLECTIONS:
        collection = "regulatory_docs"

    # Resolve relative paths from project root or cwd
    if not _os.path.isabs(folder_path):
        # Try relative to project root (where src/ is)
        project_root = _os.path.dirname(_os.path.dirname(_os.path.abspath(__file__)))
        resolved = _os.path.normpath(_os.path.join(project_root, folder_path))
        if not _os.path.isdir(resolved):
            # Try current working directory
            resolved = _os.path.normpath(_os.path.join(_os.getcwd(), folder_path))
        folder_path = resolved

    if not folder_path or not _os.path.isdir(folder_path):
        raise HTTPException(status_code=400, detail=f"文件夹不存在: {folder_path}")

    SUPPORTED = {".txt", ".md", ".docx", ".csv", ".json", ".py", ".sql", ".js", ".ts", ".java", ".go", ".sh", ".cpp", ".h", ".html", ".css", ".yaml", ".yml", ".toml", ".ini", ".cfg"}
    files_found = []
    for root, _dirs, files in _os.walk(folder_path):
        for fname in files:
            if _os.path.splitext(fname)[1].lower() in SUPPORTED:
                files_found.append(_os.path.join(root, fname))

    if not files_found:
        return {"imported": 0, "total_chunks": 0, "files_found": 0, "message": "文件夹中没有支持的文件"}

    ret = _get_retriever_for(collection)
    total_chunks = 0
    errors = []
    imported = 0

    for filepath in files_found:
        fname = _os.path.basename(filepath)
        ext = _os.path.splitext(fname)[1].lower()
        try:
            content = ""
            if ext == ".docx":
                try:
                    from docx import Document
                    doc = Document(filepath)
                    content = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
                except Exception as e:
                    errors.append(f"{fname}: {str(e)[:80]}")
                    continue
            else:
                for enc in ("utf-8", "gbk", "latin-1"):
                    try:
                        with open(filepath, "r", encoding=enc) as fh:
                            content = fh.read()
                        break
                    except UnicodeDecodeError:
                        continue

            if not content or not content.strip():
                continue

            chunks = ret.index_documents([{
                "content": content,
                "url": f"file://{filepath}",
                "metadata": {"source": "folder_import", "filename": fname, "file_type": ext},
            }])
            total_chunks += chunks
            imported += 1
        except Exception as e:
            errors.append(f"{fname}: {str(e)[:80]}")

    return {
        "imported": imported, "total_chunks": total_chunks,
        "total_in_store": ret.vector_store.count(),
        "files_found": len(files_found),
        "collection": collection,
        "errors": errors[:20],
    }


# ===================================================================
#  Search, Crawl
# ===================================================================

@router.post("/search")
async def search(request: SearchRequest):
    collection = request.filter_type if request.filter_type in KNOWN_COLLECTIONS else "regulatory_docs"
    ret = _get_retriever_for(collection)
    results = ret.retrieve(query=request.query, top_k=request.top_k)
    return {
        "query": request.query,
        "collection": collection,
        "results": [
            {"document": r.document, "metadata": r.metadata, "score": round(r.score, 4)}
            for r in results
        ],
    }


@router.post("/crawl")
async def crawl(request: CrawlRequest):
    if request.crawler_type == "doc":
        crawler = DocCrawler(max_pages=request.max_pages)
    else:
        crawler = CodeCrawler()
    docs = await crawler.crawl(request.url)
    return {
        "crawled_count": len(docs),
        "documents": [
            {"url": d.url, "title": d.title, "content": d.content,
             "source_type": d.source_type, "metadata": d.metadata, "crawled_at": d.crawled_at}
            for d in docs
        ],
    }


@router.post("/pipeline/crawl-and-index")
async def crawl_and_index(request: CrawlRequest):
    if request.crawler_type == "doc":
        crawler = DocCrawler(max_pages=request.max_pages)
    else:
        crawler = CodeCrawler()
    docs = await crawler.crawl(request.url)

    if not docs:
        return {"crawled_count": 0, "indexed_count": 0, "message": "No documents crawled"}

    # Use explicit collection if provided, else infer from crawler_type
    if request.collection in KNOWN_COLLECTIONS:
        collection = request.collection
    else:
        collection = "code_docs" if request.crawler_type == "code" else "regulatory_docs"
    ret = _get_retriever_for(collection)
    index_payload = [
        {"content": d.content, "url": d.url, "metadata": {"title": d.title, "source_type": d.source_type, **d.metadata}}
        for d in docs
    ]
    indexed = ret.index_documents(index_payload)

    return {
        "crawled_count": len(docs),
        "indexed_count": indexed,
        "total_in_store": ret.vector_store.count(),
        "collection": collection,
        "sources": [d.url for d in docs],
    }


# ===================================================================
#  Scenarios
# ===================================================================

@router.get("/scenarios")
async def list_scenarios():
    from src.business.banking import HSBC_DESCRIPTIONS, HSBC_SCHEMAS
    return {
        key: {"name": key, "description": HSBC_DESCRIPTIONS.get(key, ""),
              "table_count": HSBC_SCHEMAS[key].count("CREATE TABLE")}
        for key in HSBC_SCHEMAS
    }


@router.get("/scenarios/{name}")
async def get_scenario(name: str):
    from src.business.banking import get_schema
    ddl = get_schema(name)
    if not ddl:
        raise HTTPException(status_code=404, detail=f"Scenario '{name}' not found")
    return {"name": name, "ddl": ddl}


# ===================================================================
#  Metrics & Benchmark
# ===================================================================

@router.get("/metrics", response_model=MetricsSummary)
async def get_metrics():
    tracker = get_metrics_tracker()
    summary = tracker.get_summary()
    return MetricsSummary(**summary)


@router.post("/benchmark")
async def run_benchmark(request: BenchmarkRequest):
    runner = BenchmarkRunner(use_mock_llm=request.use_mock_llm)
    results = await runner.run()
    return results


@router.get("/history")
async def get_history(limit: int = 50):
    """获取最近的操作历史记录"""
    tracker = get_metrics_tracker()
    return {"history": tracker.get_history(limit=limit)}


@router.delete("/history")
async def clear_history():
    """清空历史记录"""
    tracker = get_metrics_tracker()
    count = len(tracker.get_history(limit=99999))
    tracker.reset()
    return {"cleared": True, "previous_count": count}


# ===================================================================
#  制度文档问答 & 文档审查
# ===================================================================

from src.api.schemas import DocQARequest, DocReviewRequest, DocClassifyRequest


@router.post("/documents/qa")
async def document_qa(request: DocQARequest):
    ret = _get_retriever_for("regulatory_docs")
    context = ret.retrieve_context(request.question, top_k=5)
    from src.llm.prompts import DOC_QA_SYSTEM, DOC_QA_USER
    from src.llm.providers import OpenAIProvider
    provider = OpenAIProvider()
    system_prompt = DOC_QA_SYSTEM.format(context=context if context else "暂无相关制度文档")
    user_prompt = DOC_QA_USER.format(question=request.question)
    response = await provider.generate([
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": user_prompt},
    ])
    tokens = response.usage.get("total_tokens", 0)
    get_metrics_tracker().record_generation(
        gen_type="qa", success=True,
        generation_time_ms=0, validation_passed=True, rag_used=bool(context),
        tokens_used=tokens,
        metadata={"question": request.question[:100]},
    )
    return {
        "question": request.question,
        "answer": response.content,
        "model": response.model,
        "sources": [
            {"document": r.document[:200], "score": round(r.score, 4), "url": r.metadata.get("url", "")}
            for r in ret.retrieve(request.question, top_k=3)
        ],
        "tokens_used": tokens,
    }


@router.post("/documents/review")
async def document_review(request: DocReviewRequest):
    ret = _get_retriever_for("regulatory_docs")
    context = ret.retrieve_context(request.content[:500], top_k=5)
    from src.llm.prompts import DOC_REVIEW_SYSTEM, DOC_REVIEW_USER
    from src.llm.providers import OpenAIProvider
    provider = OpenAIProvider()
    response = await provider.generate_with_json([
        {"role": "system", "content": DOC_REVIEW_SYSTEM.format(context=context if context else "暂无参考制度规范")},
        {"role": "user", "content": DOC_REVIEW_USER.format(title=request.title, content=request.content[:3000])},
    ])
    from src.config import settings as _settings

    if isinstance(response, dict):
        items_raw = _find_items_array(response)
        normalized_items = []
        for item in items_raw:
            normalized_items.append({
                "status": item.get("status") or item.get("compliance_status") or item.get("severity") or "未知",
                "clause": item.get("clause") or item.get("regulation") or item.get("rule") or "",
                "detail": item.get("detail") or item.get("details") or item.get("issue") or item.get("violation_detail") or "",
                "suggestion": item.get("suggestion") or item.get("recommendation") or "",
            })
        risk_level = "未知"
        for key in ("risk_level", "overall_risk_rating", "risk_rating", "overall_rating"):
            val = _find_value_recursive(response, key)
            if val and isinstance(val, str):
                risk_level = val
                break
        summary = ""
        for key in ("summary", "overall_assessment", "overall_summary", "risk_justification"):
            val = _find_value_recursive(response, key)
            if val and isinstance(val, str):
                summary = val
                break
        normalized = {"risk_level": risk_level, "items": normalized_items, "summary": summary}
    else:
        normalized = response

    source_count = len(ret.retrieve(request.content[:300], top_k=3))
    get_metrics_tracker().record_generation(
        gen_type="review", success=True,
        generation_time_ms=0, validation_passed=True,
        rag_used=source_count > 0,
        tokens_used=len(request.content) // 2,  # rough estimate
        metadata={"title": request.title[:80]},
    )
    return {
        "title": request.title,
        "review_result": normalized,
        "model": _settings.llm_model,
        "sources_count": source_count,
    }


@router.post("/documents/classify")
async def document_classify(request: DocClassifyRequest):
    from src.llm.prompts import DOC_CLASSIFY_SYSTEM
    from src.llm.providers import OpenAIProvider
    provider = OpenAIProvider()
    response = await provider.generate_with_json([
        {"role": "system", "content": "You are a banking document classification expert. Always return valid JSON."},
        {"role": "user", "content": DOC_CLASSIFY_SYSTEM.format(content=request.content[:5000])},
    ])
    if response.get("parse_error") or "raw" in response:
        return {
            "title": request.title,
            "classification": {
                "doc_type": "未知", "category": "未分类", "risk_level": "未知",
                "summary": "AI 分类解析失败", "effective_date": "", "keywords": [], "key_clauses": [],
                "raw_response": response.get("raw", str(response)),
            },
        }
    get_metrics_tracker().record_generation(
        gen_type="classify", success=True,
        generation_time_ms=0, validation_passed=True, rag_used=False,
        tokens_used=len(request.content) // 2,
        metadata={"title": request.title[:80]},
    )
    return {"title": request.title, "classification": response}


@router.get("/dashboard/status")
async def dashboard_status():
    stats = get_metrics_tracker().get_summary()
    reg_ret = _get_retriever_for("regulatory_docs")
    code_ret = _get_retriever_for("code_docs")
    return {
        "service": "rag-code-gen",
        "regulatory_docs_count": reg_ret.vector_store.count(),
        "code_docs_count": code_ret.vector_store.count(),
        "vector_store_count": reg_ret.vector_store.count() + code_ret.vector_store.count(),
        "sql_generated": stats["sql"]["total_generations"],
        "python_generated": stats["python"]["total_generations"],
        "total_tokens": stats["total_tokens_used"],
        "sql_success_rate": stats["sql"]["success_rate"],
        "python_success_rate": stats["python"]["success_rate"],
    }
